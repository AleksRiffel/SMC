from docxtpl import DocxTemplate
import openpyxl
from openpyxl import Workbook
from smc_package.schemas.GeneralDict import *
from docx import Document
from docx.document import Document as DocumentObject

def load_excel(excel_file_path: str) -> Workbook:
    """
    Загружает Excel-файл с данными.

    Args:
        excel_file_path (str): Путь к Excel-файлу.

    Returns:
        Workbook: Объект рабочей книги Excel.
    """
    return openpyxl.load_workbook(filename=excel_file_path)

def load_OPOP(opop_file_path:str)-> DocumentObject:
    """
    Загружает ОПОП и возвращает в нужном формате
    Args:
        opop_file_path(str): путь до ОПОП
    Returns:
        DocumentObject: Объект word-документа
    """
    return Document(opop_file_path)


def extract_basic_info(wb: Workbook) -> dict:
    """
    Извлекает общую информацию о направлении, профиле и списке дисциплин.

    Args:
        wb (Workbook): Объект рабочей книги Excel.

    Returns:
        dict: Словарь с ключами:
            - "напрапвление": Название направления.
            - "профиль": Название профиля.
            - "список дисциплин": Список дисциплин.
            - "кафедры": Список кафедр.
    """
    names = []
    names.append(wb['ПланСвод']["C6"].value)
    for i in range(7, 10000):
        if wb['ПланСвод']["C" + str(i)].value is None:
            if wb['ПланСвод']["C" + str(i + 1)].value is None:
                break
            continue
        if wb['ПланСвод']["C" + str(i - 1)].value is None or wb['ПланСвод']["C" + str(i)].font.bold:
            continue
        names.append(wb['ПланСвод']["C" + str(i)].value)
    namesKafedras = []
    for i in range(2, 10000):
        if wb['Кафедры']["C" + str(i)].value is None:
            break
        namesKafedras.append(wb['Кафедры']["C" + str(i)].value)
    return {
    "направление":str(wb['Титул']["B18"].value.split("Направление: ")[1]),
    "профиль": wb['Титул']["B19"].value,
    "список дисциплин": names,
    "кафедры": namesKafedras

    }

def extract_specific_info(wb: Workbook) -> dict:
    """
    Извлекает специфичную информацию из файла (например, форму обучения).

    Args:
        wb (Workbook): Объект рабочей книги Excel.

    Returns:
        dict: Словарь с ключами:
            - "форма обучения": Форма обучения.
    """
    return {
        "форма обучения": wb['Титул']["A31"].value.split("Форма обучения: ")[1].split(" ")[0].lower()
    }


def extract_competences(wb: Workbook, pickedname: str) -> list:
    """
    Извлекает коды компетенций для указанной дисциплины.

    Args:
        wb (Workbook): Объект рабочей книги Excel.
        pickedname (str): Название дисциплины.

    Returns:
        list: Список кодов компетенций.
    """
    codesCompitence = []
    for i in range(4, 10000):
        if wb['Компетенции(2)']["F" + str(i)].value is None:
            break
        if wb['Компетенции(2)']["F" + str(i)].value == pickedname:
            codesCompitence = wb['Компетенции(2)']["G" + str(i)].value.split("; ")
            break
    return codesCompitence


def fill_competences(wb: Workbook, codesCompitence: list) -> list:
    """
    Заполняет компетенции в GeneralDict на основе предоставленных кодов.

    Args:
        wb (Workbook): Объект рабочей книги Excel.
        codesCompitence (list): Список кодов компетенций.

    Returns:
        list: Список объектов Компетенция.
    """
    compitences = []
    for i in range(3, 10000):
        if wb['Компетенции']["D" + str(i)].value is None:
            break
        if wb['Компетенции']["B" + str(i)].value is None:
            continue
        for j in codesCompitence:
            if wb['Компетенции']["B" + str(i)].value == j:
                compitences.append(Компетенция(
                    код=j,
                    описание=wb['Компетенции']["D" + str(i)].value
                ))
                break
        if len(compitences) == len(codesCompitence):
            break
    return compitences


def fill_discipline_volume(wb: Workbook, pickedname: str) -> dict:
    """
    Заполняет объем дисциплины: часы, зачетные единицы, виды занятий.

    Примечание:
        Функция требует доработки. В текущей реализации работает только для дисциплины, 
        изуч��емой в одном семестре.

    Args:
        wb (Workbook): Объект рабочей книги Excel.
        pickedname (str): Название дисциплины.

    Returns:
        dict: Словарь с ключами:
            - "часы": Общее количество часов.
            - "зачетные единицы": Количество зачетных единиц.
            - "виды занятий": Список объектов ХарактеристикиОбъемаДисциплины.
    """
    volume=[]
    часы=""
    зачетные_единицы=""
    for i in range(7, 10000):
        if wb['План']["C" + str(i)].value == pickedname:
            часы = wb['План']["L" + str(i)].value
            зачетные_единицы = wb['План']["I" + str(i)].value
            #Определение семестров в предмете 0-индекс тип аттестации, 1-индекс номер семестра
            listsemestrs=[]
            if wb['План']["D" + str(i)].value != None:
                for number in list(str(wb['План']["D" + str(i)].value)):
                    listsemestrs.append(["экзамен",int(number)])
            if wb['План']["E" + str(i)].value != None:
                for number in list(str(wb['План']["E" + str(i)].value)):
                    listsemestrs.append(["экзамен",int(number)])
            listsemestrs.sort(key=lambda x: x[1])
            for infosemestr in listsemestrs:
                infosemestr[1]=infosemestr[1]
                hoursdict=ХарактеристикиОбъемаДисциплины(
                лекции=str(wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+1).value) if wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+1).value!=None else "-",
                практика=str(wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+3).value) if wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+3).value!=None else "-",
                лабраб=str(wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+2).value) if wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+2).value!=None else "-", 
                самраб=str(wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+5).value) if wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+5).value!=None else "-",
                контроль=str(wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+6).value) if wb['План'].cell(row=i,column=17+(infosemestr[1]-1)*7+6).value!=None else "-",
                аттестация=infosemestr[0],
                семестр=str(infosemestr[1]),
                курс=str(infosemestr[1]//2+infosemestr[1]%2)
                )
                volume.append(hoursdict)
            break
    return {
        "часы":часы,
        "зачетные единицы":зачетные_единицы,
        "виды занятий":volume
    }


def check_coursework(wb: Workbook, pickedname: str) -> bool:
    """
    Проверяет наличие курсовой работы для указанной дисциплины.

    Примечание:
        Функция требует доработки. В текущей реализации отсутствует функционал
        определения семестра, в котором проводится курсовая работа.

    Args:
        wb (Workbook): Объект рабочей книги Excel.
        pickedname (str): Название дисциплины.

    Returns:
        bool: True, если курсовая работа есть, иначе False.
    """
    isworks = False
    for i in range(2, 1000):
        if wb['Курсовые']["A" + str(i)].value == pickedname + " ":
            isworks = True
            break
    return isworks

def fill_IDK_from_OPOP(compitences:list[Компетенция],opop:DocumentObject)-> list[Компетенция]:
    """
    Заполняет блок компитенций ИДК

    Args:
        compitence(list[Компитенция]):список объектов, описывающих компитенции. Заполнен код компитенции для его обработки. ОБЯЗАТЕЛЬНО
        opop(DocumentObject): документ ОПОП
    Returns:
        list[Компитенция]: заполненные данными список компитенций и их ИДК
    """
    doc = opop
    previous_cell=""
    # Перебор всех таблиц в документе
    indx_current_compitention=0
    for table_index, table in enumerate(doc.tables):
        # Нахождение таблицы с УК и ОПК. Для поиска определяется первая ячейка первой строки и составляется строка без пробелов строчными для точного определения таблиц
        if ''.join(''.join(paragraph.text for paragraph in table.rows[0].cells[0].paragraphs).split(" ")).lower()=="наименованиекатегории(группы)ук" or ''.join(''.join(paragraph.text for paragraph in table.rows[0].cells[0].paragraphs).split(" ")).lower()=="наименованиекатегории(группы)опк":
            # определение необходимости первого ИДК
            skipGroup=True
            previous_cell=(''.join(paragraph.text for paragraph in table.rows[1].cells[1].paragraphs)).split(".")[0]
            for indx_compitence in range(indx_current_compitention,len(compitences)):
                    if "".join((compitences[indx_compitence].код).lower().split(" "))==previous_cell.lower():
                        print("match",previous_cell.lower())
                        skipGroup=False
                        indx_current_compitention=indx_compitence
                        break
            # поиск ИДК с их группировкой по компитенции с дальнейшей записью в модель
            for row in table.rows[1:]: #начинаем с первой строки для пропуска хедера
                #если ИДК относится к текущей компитенции
                if (' '.join(paragraph.text for paragraph in row.cells[1].paragraphs)).split(".")[0]==previous_cell:
                    if skipGroup:
                        continue
                    compitences[indx_current_compitention].идк.append(ИДК(
                        описание=' '.join(paragraph.text for paragraph in row.cells[2].paragraphs))
                    )
                    continue

                #проверка текущей группы на соответствие компитенции
                skipGroup=True
                previous_cell=(''.join(paragraph.text for paragraph in row.cells[1].paragraphs)).split(".")[0]
                # поиск подходящей компитенции в модели
                for compitence in range(indx_current_compitention,len(compitences)):
                    if "".join((compitences[compitence].код).lower().split(" "))==previous_cell.lower():
                        print("match",previous_cell.lower())
                        indx_current_compitention=compitence
                        skipGroup=False
                        break
                # проверка для добавления ИДК и добаления первой ИДК новой группы
                if skipGroup:
                    continue
                compitences[indx_current_compitention].идк.append(ИДК(
                        описание=' '.join(paragraph.text for paragraph in row.cells[2].paragraphs))
                    )
    return compitences

def generate_documents(
        content:GeneralDictModel,
        template_rp: str = "шаблонРП.docx",
        template_a: str = "шаблонА.docx",
        template_fos: str ="шаблонФОС.dox",
        output_rp: str = "generated_RP.docx", 
        output_a: str = "generated_A.docx",
        output_fos:str = "generated_fos.docx") -> None:
    """
    Генерирует документы по шаблонам на основе предоставленного контента.

    Args:
        content (GeneralDictModel): Данные для заполнения шаблонов.
        template_rp (str, optional): Путь к шаблону РП. Defaults to "шаблонРП.docx".
        template_a (str, optional): Путь к шаблону А. Defaults to "шаблонА.docx".
        template_fos (str, optional): Путь к шаблону ФОС. Defaults to "шаблонФОС.dox".
        output_rp (str, optional): Путь для сохранения РП. Defaults to "generated_RP.docx".
        output_a (str, optional): Путь для сохранения А. Defaults to "generated_A.docx".
        output_fos (str, optional): Путь для сохранения ФОС. Defaults to "generated_fos.docx".
    """
    doc = DocxTemplate(template_rp)
    doc.render(content.model_dump())
    doc.save(output_rp)
    doc = DocxTemplate(template_a)
    doc.render(content.model_dump())
    doc.save(output_a)